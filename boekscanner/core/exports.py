"""Export-module: doorzoekbare PDF, DOCX, TXT, MD.

Per pagina én als één gecombineerd bestand. Maakt gebruik van:

- ``ocrmypdf`` → doorzoekbare PDF (origineel beeld + onzichtbare tekstlaag)
- ``python-docx`` → Word-document met tekst en (optioneel) ingesloten beeld
- ``pypdf`` → samenvoegen van per-pagina PDF's
- Native string-write voor TXT en Markdown
"""

from __future__ import annotations

import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import List, Literal, Optional

from docx import Document
from docx.shared import Inches, Pt
from loguru import logger
from pypdf import PdfReader, PdfWriter

from .config import get_config
from .image_proc import load_image, save_image
from .projects import PageMeta, Project


ExportFormat = Literal["pdf", "docx", "txt", "md"]


@dataclass
class ExportResult:
    format: ExportFormat
    combined_path: Optional[Path]
    per_page_paths: List[Path]


# --- Helpers ---------------------------------------------------------------


def _processed_or_raw(project: Project, page: PageMeta) -> Path:
    """Geef het beste beschikbare beeld voor een pagina (processed > raw)."""
    if page.processed_filename:
        p = project.processed_dir / page.processed_filename
        if p.is_file():
            return p
    return project.raw_dir / page.raw_filename


def _read_ocr_text(project: Project, page: PageMeta) -> str:
    if not page.ocr_filename:
        return ""
    p = project.ocr_dir / page.ocr_filename
    return p.read_text(encoding="utf-8") if p.is_file() else ""


def _safe_export_dir(project: Project, fmt: ExportFormat) -> Path:
    out = project.export_dir / fmt
    out.mkdir(parents=True, exist_ok=True)
    return out


# --- TXT / Markdown --------------------------------------------------------


def export_txt(project: Project, *, per_page: bool = True, combined: bool = True) -> ExportResult:
    out_dir = _safe_export_dir(project, "txt")
    per_paths: List[Path] = []
    combined_buf: List[str] = []
    for idx, page in enumerate(project.meta.pages, start=1):
        text = _read_ocr_text(project, page)
        if per_page:
            p = out_dir / f"pagina-{idx:03d}.txt"
            p.write_text(text, encoding="utf-8")
            per_paths.append(p)
        combined_buf.append(f"--- Pagina {idx} ---\n{text}\n")
    combined_path: Optional[Path] = None
    if combined:
        combined_path = out_dir / f"{project.meta.slug}-volledig.txt"
        combined_path.write_text("\n".join(combined_buf), encoding="utf-8")
    return ExportResult(format="txt", combined_path=combined_path, per_page_paths=per_paths)


def export_md(project: Project, *, per_page: bool = True, combined: bool = True) -> ExportResult:
    out_dir = _safe_export_dir(project, "md")
    per_paths: List[Path] = []
    combined_buf: List[str] = [f"# {project.meta.name}\n"]
    if project.meta.description:
        combined_buf.append(f"_{project.meta.description}_\n")
    for idx, page in enumerate(project.meta.pages, start=1):
        text = _read_ocr_text(project, page)
        if per_page:
            md = f"# {project.meta.name} — Pagina {idx}\n\n{text}\n"
            p = out_dir / f"pagina-{idx:03d}.md"
            p.write_text(md, encoding="utf-8")
            per_paths.append(p)
        combined_buf.append(f"\n## Pagina {idx}\n\n{text}\n")
    combined_path: Optional[Path] = None
    if combined:
        combined_path = out_dir / f"{project.meta.slug}-volledig.md"
        combined_path.write_text("\n".join(combined_buf), encoding="utf-8")
    return ExportResult(format="md", combined_path=combined_path, per_page_paths=per_paths)


# --- DOCX ------------------------------------------------------------------


def _add_page_to_docx(doc: Document, page_number: int, text: str, image_path: Optional[Path]) -> None:
    heading = doc.add_heading(f"Pagina {page_number}", level=2)
    heading.runs[0].font.size = Pt(14)
    if image_path and image_path.is_file() and image_path.suffix.lower() in {
        ".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tif", ".tiff",
    }:
        try:
            doc.add_picture(str(image_path), width=Inches(5.5))
        except Exception as exc:
            logger.warning("Kan beeld niet invoegen in DOCX ({}): {}", image_path, exc)
    para = doc.add_paragraph(text)
    for run in para.runs:
        run.font.size = Pt(11)


def export_docx(
    project: Project,
    *,
    per_page: bool = True,
    combined: bool = True,
    include_images: bool = True,
) -> ExportResult:
    out_dir = _safe_export_dir(project, "docx")
    per_paths: List[Path] = []

    if per_page:
        for idx, page in enumerate(project.meta.pages, start=1):
            text = _read_ocr_text(project, page)
            doc = Document()
            doc.add_heading(f"{project.meta.name} — Pagina {idx}", level=1)
            img = _processed_or_raw(project, page) if include_images else None
            _add_page_to_docx(doc, idx, text, img)
            p = out_dir / f"pagina-{idx:03d}.docx"
            doc.save(str(p))
            per_paths.append(p)

    combined_path: Optional[Path] = None
    if combined:
        doc = Document()
        doc.add_heading(project.meta.name, level=0)
        if project.meta.description:
            doc.add_paragraph(project.meta.description)
        doc.add_paragraph(f"Geëxporteerd op {datetime.now().strftime('%d-%m-%Y %H:%M')}")
        for idx, page in enumerate(project.meta.pages, start=1):
            text = _read_ocr_text(project, page)
            img = _processed_or_raw(project, page) if include_images else None
            _add_page_to_docx(doc, idx, text, img)
            doc.add_page_break()
        combined_path = out_dir / f"{project.meta.slug}-volledig.docx"
        doc.save(str(combined_path))

    return ExportResult(format="docx", combined_path=combined_path, per_page_paths=per_paths)


# --- Doorzoekbare PDF (ocrmypdf) ------------------------------------------


def _is_ocrmypdf_available() -> bool:
    try:
        import ocrmypdf  # noqa: F401
        return True
    except Exception:
        return False


def _image_to_searchable_pdf(image: Path, output: Path, languages: List[str]) -> Path:
    """Maak één doorzoekbare PDF van één afbeelding."""
    import ocrmypdf  # lazy import (laden duurt even)

    cfg = get_config()
    # ocrmypdf accepteert beelden als input wanneer --image-dpi gezet is
    ocrmypdf.ocr(
        input_file=str(image),
        output_file=str(output),
        language="+".join(languages),
        image_dpi=300,
        deskew=False,            # we deskewen al zelf
        rotate_pages=False,
        clean=False,
        optimize=1,
        progress_bar=False,
    )
    return output


def _pdf_to_searchable_pdf(input_pdf: Path, output: Path, languages: List[str]) -> Path:
    import ocrmypdf

    ocrmypdf.ocr(
        input_file=str(input_pdf),
        output_file=str(output),
        language="+".join(languages),
        skip_text=False,
        force_ocr=True,
        deskew=False,
        rotate_pages=False,
        progress_bar=False,
    )
    return output


def export_searchable_pdf(
    project: Project,
    *,
    per_page: bool = True,
    combined: bool = True,
) -> ExportResult:
    """Genereer doorzoekbare PDF's. Vereist ``ocrmypdf`` + Tesseract + Ghostscript."""
    if not _is_ocrmypdf_available():
        raise RuntimeError(
            "ocrmypdf is niet geïnstalleerd. Voer `pip install ocrmypdf` uit "
            "en zorg dat Ghostscript geïnstalleerd staat."
        )

    out_dir = _safe_export_dir(project, "pdf")
    languages = project.meta.settings.languages or ["eng"]
    per_paths: List[Path] = []

    for idx, page in enumerate(project.meta.pages, start=1):
        src = _processed_or_raw(project, page)
        target = out_dir / f"pagina-{idx:03d}.pdf"
        try:
            if src.suffix.lower() == ".pdf":
                _pdf_to_searchable_pdf(src, target, languages)
            else:
                _image_to_searchable_pdf(src, target, languages)
        except Exception as exc:
            logger.error("Kon doorzoekbare PDF voor pagina {} niet maken: {}", idx, exc)
            continue
        per_paths.append(target)

    combined_path: Optional[Path] = None
    if combined and per_paths:
        combined_path = out_dir / f"{project.meta.slug}-volledig.pdf"
        writer = PdfWriter()
        for p in per_paths:
            try:
                reader = PdfReader(str(p))
                for pg in reader.pages:
                    writer.add_page(pg)
            except Exception as exc:
                logger.warning("Kan {} niet samenvoegen: {}", p, exc)
        with combined_path.open("wb") as f:
            writer.write(f)

    if not per_page:
        for p in per_paths:
            try:
                p.unlink()
            except OSError:
                pass
        per_paths = []

    return ExportResult(format="pdf", combined_path=combined_path, per_page_paths=per_paths)


# --- Generieke entry-point -------------------------------------------------


def export_project(
    project: Project,
    formats: List[ExportFormat],
    *,
    per_page: bool = True,
    combined: bool = True,
) -> List[ExportResult]:
    results: List[ExportResult] = []
    for fmt in formats:
        if fmt == "txt":
            results.append(export_txt(project, per_page=per_page, combined=combined))
        elif fmt == "md":
            results.append(export_md(project, per_page=per_page, combined=combined))
        elif fmt == "docx":
            results.append(export_docx(project, per_page=per_page, combined=combined))
        elif fmt == "pdf":
            results.append(export_searchable_pdf(project, per_page=per_page, combined=combined))
        else:
            raise ValueError(f"Onbekend export-formaat: {fmt}")
    return results
